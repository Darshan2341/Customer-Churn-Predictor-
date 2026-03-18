import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go
from sklearn.ensemble import RandomForestClassifier, GradientBoostingClassifier
from sklearn.linear_model import LogisticRegression
from sklearn.model_selection import train_test_split
from sklearn.metrics import accuracy_score, confusion_matrix, roc_auc_score, roc_curve, classification_report
from sklearn.preprocessing import LabelEncoder, StandardScaler
from io import BytesIO
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import warnings
warnings.filterwarnings('ignore')

st.set_page_config(page_title="Churn Predictor", page_icon="🔮", layout="wide")

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Syne:wght@600;700;800&family=Inter:wght@300;400;500&display=swap');
html, body, [class*="css"] { font-family: 'Inter', sans-serif; }
.main { background: #07090f; }
.block-container { padding: 2rem 3rem; }

.hero { margin-bottom: 2rem; }
.hero h1 {
    font-family: 'Syne', sans-serif;
    font-size: 2.6rem; font-weight: 800;
    background: linear-gradient(120deg, #38bdf8, #818cf8, #34d399);
    -webkit-background-clip: text; -webkit-text-fill-color: transparent;
    margin-bottom: 0.2rem;
}
.hero p { color: #475569; font-size: 0.95rem; font-weight: 300; }

.kpi {
    background: #0d1117; border: 1px solid #1e293b;
    border-radius: 14px; padding: 1.2rem 1.4rem;
    border-top: 2px solid #38bdf8;
}
.kpi-val { font-family: 'Syne', sans-serif; font-size: 2rem; font-weight: 700; color: #e2e8f0; }
.kpi-lbl { color: #475569; font-size: 0.72rem; text-transform: uppercase; letter-spacing: 1.5px; margin-top: 2px; }
.kpi-tag { font-size: 0.72rem; padding: 2px 8px; border-radius: 20px; margin-top: 6px; display: inline-block; }
.tag-g { background: rgba(52,211,153,0.1); color: #34d399; }
.tag-r { background: rgba(248,113,113,0.1); color: #f87171; }
.tag-b { background: rgba(56,189,248,0.1); color: #38bdf8; }

.sec { font-family: 'Syne', sans-serif; font-size: 1rem; font-weight: 700;
    color: #94a3b8; border-bottom: 1px solid #1e293b;
    padding-bottom: 0.4rem; margin: 1.5rem 0 1rem; }

.card {
    background: #0d1117; border: 1px solid #1e293b;
    border-left: 3px solid #38bdf8; border-radius: 10px;
    padding: 0.9rem 1.1rem; margin-bottom: 0.6rem;
    color: #cbd5e1; font-size: 0.88rem; line-height: 1.6;
}
.card-g { border-left-color: #34d399; }
.card-r { border-left-color: #f87171; }
.card-y { border-left-color: #fbbf24; }

.stButton > button {
    background: linear-gradient(135deg, #0ea5e9, #6366f1) !important;
    color: white !important; border: none !important;
    border-radius: 10px !important; font-family: 'Syne', sans-serif !important;
    font-weight: 700 !important; font-size: 0.95rem !important;
    padding: 0.65rem 1rem !important; width: 100% !important;
}
.stDownloadButton > button {
    background: #0d1117 !important; color: #38bdf8 !important;
    border: 1px solid #1e3a5f !important; border-radius: 10px !important;
    font-family: 'Syne', sans-serif !important; font-weight: 600 !important;
    width: 100% !important;
}
section[data-testid="stSidebar"] { background: #090c14 !important; border-right: 1px solid #1e293b !important; }
</style>
""", unsafe_allow_html=True)


# ── Helpers ───────────────────────────────────────────────────

@st.cache_data
def sample_data(n=600):
    np.random.seed(42)
    return pd.DataFrame({
        'Age': np.random.randint(18, 75, n),
        'Gender': np.random.choice(['Male', 'Female'], n),
        'Tenure': np.random.randint(1, 72, n),
        'MonthlyCharges': np.round(np.random.uniform(20, 120, n), 2),
        'TotalCharges': np.round(np.random.uniform(100, 8000, n), 2),
        'Contract': np.random.choice(['Month-to-month', 'One year', 'Two year'], n, p=[0.5, 0.3, 0.2]),
        'InternetService': np.random.choice(['DSL', 'Fiber optic', 'No'], n),
        'TechSupport': np.random.choice(['Yes', 'No'], n),
        'PaymentMethod': np.random.choice(['Electronic check', 'Mailed check', 'Bank transfer', 'Credit card'], n),
        'NumSupportCalls': np.random.randint(0, 10, n),
        'Churn': np.random.choice(['Yes', 'No'], n, p=[0.27, 0.73])
    })


def smart_read(file):
    for sep in [',', ';', '\t', '|']:
        try:
            df = pd.read_csv(file, sep=sep, engine='python')
            if df.shape[1] > 1:
                return df
            file.seek(0)
        except:
            try: file.seek(0)
            except: pass
    return pd.read_csv(file)


def preprocess(df, target):
    df = df.copy().dropna()
    drops = [c for c in df.columns if 'id' in c.lower() or 'name' in c.lower()]
    df.drop(columns=drops, inplace=True, errors='ignore')
    le = LabelEncoder()
    for col in df.select_dtypes(include='object').columns:
        if col != target:
            df[col] = le.fit_transform(df[col].astype(str))
    if df[target].dtype == object:
        df[target] = le.fit_transform(df[target].astype(str))
    return df


@st.cache_data
def run_model(df_hash, target, model_name):
    return None  # placeholder for cache key


def train_model(df, target, model_name):
    df = preprocess(df, target)
    X = df.drop(columns=[target])
    y = df[target]
    X_train, X_test, y_train, y_test = train_test_split(
        X, y, test_size=0.2, random_state=42, stratify=y)
    scaler = StandardScaler()
    Xtr = scaler.fit_transform(X_train)
    Xte = scaler.transform(X_test)

    mdl_map = {
        'Random Forest': RandomForestClassifier(n_estimators=150, random_state=42, n_jobs=-1),
        'Gradient Boosting': GradientBoostingClassifier(n_estimators=100, random_state=42),
        'Logistic Regression': LogisticRegression(max_iter=1000, random_state=42,)
    }
    model = mdl_map[model_name]
    model.fit(Xtr, y_train)
    y_pred = model.predict(Xte)
    y_prob = model.predict_proba(Xte)
    n_cls = len(np.unique(y_test))

    acc = accuracy_score(y_test, y_pred)
    if n_cls == 2:
        auc = roc_auc_score(y_test, y_prob[:, 1])
        fpr, tpr, _ = roc_curve(y_test, y_prob[:, 1])
    else:
        auc = roc_auc_score(y_test, y_prob, multi_class='ovr', average='macro')
        fpr, tpr = np.array([0, 1]), np.array([0, 1])

    cm = confusion_matrix(y_test, y_pred)
    if cm.shape == (2, 2):
        tn, fp, fn, tp = cm.ravel()
        prec = tp/(tp+fp) if (tp+fp) > 0 else 0
        rec  = tp/(tp+fn) if (tp+fn) > 0 else 0
        f1   = 2*prec*rec/(prec+rec) if (prec+rec) > 0 else 0
    else:
        prec = rec = f1 = acc

    fi = pd.DataFrame({'Feature': X.columns,
                        'Importance': model.feature_importances_ if hasattr(model, 'feature_importances_')
                        else np.abs(model.coef_[0])}).sort_values('Importance', ascending=False)

    return dict(model=model, scaler=scaler, acc=acc, auc=auc, cm=cm,
                fpr=fpr, tpr=tpr, fi=fi, X_cols=X.columns.tolist(),
                y_test=y_test, y_pred=y_pred, prec=prec, rec=rec, f1=f1, n_cls=n_cls)


# ── Export Excel ──────────────────────────────────────────────

def to_excel(df, r, fname):
    out = BytesIO()
    wb = openpyxl.Workbook()
    hf = Font(bold=True, color="FFFFFF", name="Calibri", size=11)
    hfill = PatternFill("solid", fgColor="0F172A")
    ctr = Alignment(horizontal="center", vertical="center")
    bdr = Border(left=Side(style='thin', color='334155'),
                 right=Side(style='thin', color='334155'),
                 top=Side(style='thin', color='334155'),
                 bottom=Side(style='thin', color='334155'))

    def style_hdr(ws):
        for cell in ws[1]:
            cell.font = hf; cell.fill = hfill
            cell.alignment = ctr; cell.border = bdr

    # Sheet 1 — Raw Data
    ws1 = wb.active; ws1.title = "📋 Raw Data"
    for ci, col in enumerate(df.columns, 1):
        ws1.cell(1, ci, col)
    for ri, row in enumerate(df.itertuples(index=False), 2):
        for ci, v in enumerate(row, 1):
            c = ws1.cell(ri, ci, v); c.border = bdr
            if ri % 2 == 0: c.fill = PatternFill("solid", fgColor="0D1117")
    style_hdr(ws1)
    for col in ws1.columns:
        ws1.column_dimensions[col[0].column_letter].width = 16

    # Sheet 2 — Statistics
    ws2 = wb.create_sheet("📊 Statistics")
    stats = df.describe().reset_index()
    for ci, col in enumerate(stats.columns, 1): ws2.cell(1, ci, col)
    for ri, row in enumerate(stats.itertuples(index=False), 2):
        for ci, v in enumerate(row, 1):
            c = ws2.cell(ri, ci, round(float(v), 4) if isinstance(v, float) else v)
            c.border = bdr
    style_hdr(ws2)
    for col in ws2.columns:
        ws2.column_dimensions[col[0].column_letter].width = 18

    # Sheet 3 — Model Performance
    ws3 = wb.create_sheet("🤖 Model Performance")
    perf = [("Metric", "Value", "Rating"),
            ("Accuracy", f"{r['acc']*100:.2f}%", "✅ Good" if r['acc'] > 0.75 else "⚠️ Improve"),
            ("ROC-AUC",  f"{r['auc']:.4f}",      "✅ Good" if r['auc'] > 0.70 else "⚠️ Improve"),
            ("Precision",f"{r['prec']*100:.2f}%", "✅ Good" if r['prec'] > 0.70 else "⚠️ Improve"),
            ("Recall",   f"{r['rec']*100:.2f}%",  "✅ Good" if r['rec'] > 0.60 else "⚠️ Improve"),
            ("F1 Score", f"{r['f1']:.4f}",        "✅ Good" if r['f1'] > 0.65 else "⚠️ Improve")]
    for ri, row in enumerate(perf, 1):
        for ci, v in enumerate(row, 1):
            c = ws3.cell(ri, ci, v); c.border = bdr; c.alignment = ctr
            if ri == 1: c.font = hf; c.fill = hfill
            elif ci == 2: c.font = Font(bold=True, color="38BDF8")
    for col in ws3.columns:
        ws3.column_dimensions[col[0].column_letter].width = 22

    # Sheet 4 — Feature Importance
    ws4 = wb.create_sheet("🏆 Feature Importance")
    for ci, h in enumerate(["Rank", "Feature", "Importance Score"], 1):
        ws4.cell(1, ci, h)
    for ri, (_, row) in enumerate(r['fi'].iterrows(), 2):
        ws4.cell(ri, 1, ri-1).border = bdr
        ws4.cell(ri, 2, row['Feature']).border = bdr
        ws4.cell(ri, 3, round(row['Importance'], 6)).border = bdr
    style_hdr(ws4)
    for col in ws4.columns:
        ws4.column_dimensions[col[0].column_letter].width = 22

    # Sheet 5 — Predictions
    ws5 = wb.create_sheet("🎯 Predictions")
    for ci, h in enumerate(["Actual", "Predicted", "Correct?"], 1):
        ws5.cell(1, ci, h)
    for ri, (a, p) in enumerate(zip(r['y_test'], r['y_pred']), 2):
        correct = "✓ Yes" if a == p else "✗ No"
        ws5.cell(ri, 1, int(a)).border = bdr
        ws5.cell(ri, 2, int(p)).border = bdr
        c = ws5.cell(ri, 3, correct); c.border = bdr
        c.font = Font(color="00C851" if a == p else "FF4444")
    style_hdr(ws5)
    for col in ws5.columns:
        ws5.column_dimensions[col[0].column_letter].width = 18

    wb.save(out)
    return out.getvalue()


# ── Export Word ───────────────────────────────────────────────

def to_word(df, r, fname):
    doc = Document()

    title = doc.add_heading('Customer Churn Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0x0E, 0xA5, 0xE9)
    title.runs[0].font.size = Pt(22)

    doc.add_paragraph(f'File: {fname}  |  Rows: {df.shape[0]:,}  |  Features: {df.shape[1]}')
    doc.add_paragraph('')

    doc.add_heading('1. Model Performance', level=1)
    t = doc.add_table(rows=1, cols=3); t.style = 'Table Grid'
    for i, h in enumerate(['Metric', 'Score', 'Status']):
        t.rows[0].cells[i].text = h
        t.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
    for m, v, s in [
        ('Accuracy',  f"{r['acc']*100:.2f}%", '✅ Good' if r['acc'] > 0.75 else '⚠️ Needs work'),
        ('ROC-AUC',   f"{r['auc']:.4f}",      '✅ Good' if r['auc'] > 0.70 else '⚠️ Needs work'),
        ('Precision', f"{r['prec']*100:.2f}%", '✅ Good' if r['prec'] > 0.70 else '⚠️ Check'),
        ('Recall',    f"{r['rec']*100:.2f}%",  '✅ Good' if r['rec'] > 0.60 else '⚠️ Check'),
        ('F1 Score',  f"{r['f1']:.4f}",        '✅ Good' if r['f1'] > 0.65 else '⚠️ Check'),
    ]:
        row = t.add_row().cells
        row[0].text = m; row[1].text = v; row[2].text = s

    doc.add_paragraph('')
    doc.add_heading('2. Top 10 Important Features', level=1)
    t2 = doc.add_table(rows=1, cols=3); t2.style = 'Table Grid'
    for i, h in enumerate(['Rank', 'Feature', 'Importance']):
        t2.rows[0].cells[i].text = h
        t2.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
    for rank, (_, row) in enumerate(r['fi'].head(10).iterrows(), 1):
        cells = t2.add_row().cells
        cells[0].text = str(rank)
        cells[1].text = str(row['Feature'])
        cells[2].text = f"{row['Importance']:.4f}"

    doc.add_paragraph('')
    doc.add_heading('3. Dataset Overview', level=1)
    doc.add_paragraph(f"Total customers: {df.shape[0]:,}")
    doc.add_paragraph(f"Total features: {df.shape[1]}")
    doc.add_paragraph(f"Missing values: {df.isnull().sum().sum()}")
    doc.add_paragraph(f"Columns: {', '.join(df.columns.tolist())}")

    doc.add_paragraph('')
    doc.add_heading('4. Summary Statistics', level=1)
    stats = df.describe().reset_index()
    t3 = doc.add_table(rows=1, cols=len(stats.columns)); t3.style = 'Table Grid'
    for i, col in enumerate(stats.columns):
        t3.rows[0].cells[i].text = str(col)
        t3.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
    for _, row in stats.iterrows():
        cells = t3.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = f"{v:.2f}" if isinstance(v, float) else str(v)

    doc.add_paragraph('')
    doc.add_heading('5. Data Sample (First 10 Rows)', level=1)
    sample = df.head(10)
    t4 = doc.add_table(rows=1, cols=len(sample.columns)); t4.style = 'Table Grid'
    for i, col in enumerate(sample.columns):
        t4.rows[0].cells[i].text = str(col)
        t4.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
    for _, row in sample.iterrows():
        cells = t4.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)

    out = BytesIO()
    doc.save(out)
    return out.getvalue()


# ── Main UI ───────────────────────────────────────────────────

st.markdown("""
<div class="hero">
  <h1>🔮 Customer Churn Predictor</h1>
  <p>Upload data · Train ML model · Get insights · Export to Excel & Word</p>
</div>
""", unsafe_allow_html=True)

# Sidebar
with st.sidebar:
    st.markdown("### ⚙️ Settings")
    model_name = st.selectbox("ML Model", ['Random Forest', 'Gradient Boosting', 'Logistic Regression'])
    use_sample = st.checkbox("Use sample data", value=True)
    st.markdown("---")
    st.markdown("### 📖 How it works")
    st.markdown("1. Upload CSV or use sample data")
    st.markdown("2. Select your Churn column")
    st.markdown("3. Click Train Model")
    st.markdown("4. Explore charts & insights")
    st.markdown("5. Export to Excel or Word")

uploaded = st.file_uploader("📂 Upload your CSV file", type=["csv"])

if uploaded:
    df = smart_read(uploaded)
    use_sample = False
    fname = uploaded.name
elif use_sample:
    df = sample_data()
    fname = "sample_data.csv"
    st.info("📊 Using built-in sample data. Check 'Use sample data' in sidebar or upload your CSV.")
else:
    st.warning("Please upload a CSV or enable sample data in the sidebar.")
    st.stop()

st.markdown("---")
default_idx = df.columns.tolist().index('Churn') if 'Churn' in df.columns else 0
target = st.selectbox("🎯 Select Churn/Target Column", df.columns.tolist(), index=default_idx)

# KPI row
st.markdown('<div class="sec">📊 Dataset Overview</div>', unsafe_allow_html=True)
num_cols = df.select_dtypes(include='number').columns.tolist()
churn_rate = round(df[target].value_counts(normalize=True).iloc[0] * 100, 1)

c1, c2, c3, c4, c5 = st.columns(5)
with c1:
    st.markdown(f'<div class="kpi"><div class="kpi-val">{len(df):,}</div><div class="kpi-lbl">Total Customers</div><div class="kpi-tag tag-b">Dataset</div></div>', unsafe_allow_html=True)
with c2:
    st.markdown(f'<div class="kpi"><div class="kpi-val">{df.shape[1]}</div><div class="kpi-lbl">Features</div><div class="kpi-tag tag-b">Columns</div></div>', unsafe_allow_html=True)
with c3:
    st.markdown(f'<div class="kpi"><div class="kpi-val">{churn_rate}%</div><div class="kpi-lbl">Churn Rate</div><div class="kpi-tag tag-r">At Risk</div></div>', unsafe_allow_html=True)
with c4:
    st.markdown(f'<div class="kpi"><div class="kpi-val">{len(num_cols)}</div><div class="kpi-lbl">Numeric Cols</div><div class="kpi-tag tag-g">ML Ready</div></div>', unsafe_allow_html=True)
with c5:
    st.markdown(f'<div class="kpi"><div class="kpi-val">{df.isnull().sum().sum()}</div><div class="kpi-lbl">Missing Values</div><div class="kpi-tag tag-g">Clean</div></div>', unsafe_allow_html=True)

st.markdown("<br>", unsafe_allow_html=True)

tab1, tab2, tab3, tab4, tab5 = st.tabs(["📋 Data", "📊 EDA", "🤖 ML Model", "🎯 Predict", "📥 Export"])

# ── Tab 1: Data ───────────────────────────────────────────────
with tab1:
    st.markdown('<div class="sec">Raw Data</div>', unsafe_allow_html=True)
    st.dataframe(df, use_container_width=True, height=400)
    st.markdown('<div class="sec">Summary Statistics</div>', unsafe_allow_html=True)
    st.dataframe(df.describe(), use_container_width=True)

# ── Tab 2: EDA ────────────────────────────────────────────────
with tab2:
    st.markdown('<div class="sec">Exploratory Data Analysis</div>', unsafe_allow_html=True)
    cat_cols = df.select_dtypes(include='object').columns.tolist()

    col_a, col_b = st.columns(2)
    with col_a:
        fig = px.pie(df, names=target, title="Churn Distribution",
                     color_discrete_sequence=['#38bdf8', '#f87171'], template='plotly_dark',
                     hole=0.4)
        fig.update_layout(paper_bgcolor='#0d1117', plot_bgcolor='#0d1117')
        st.plotly_chart(fig, use_container_width=True)
    with col_b:
        if num_cols:
            sel = st.selectbox("Numeric feature", num_cols, key='eda1')
            fig2 = px.histogram(df, x=sel, color=target, barmode='overlay',
                                title=f"{sel} by Churn", template='plotly_dark',
                                color_discrete_sequence=['#38bdf8', '#f87171'])
            fig2.update_layout(paper_bgcolor='#0d1117', plot_bgcolor='#0d1117')
            st.plotly_chart(fig2, use_container_width=True)

    cats_clean = [c for c in cat_cols if c != target]
    if cats_clean:
        col_c, col_d = st.columns(2)
        with col_c:
            sel2 = st.selectbox("Categorical feature", cats_clean, key='eda2')
            fig3 = px.histogram(df, x=sel2, color=target, barmode='group',
                                title=f"Churn by {sel2}", template='plotly_dark',
                                color_discrete_sequence=['#38bdf8', '#f87171'])
            fig3.update_layout(paper_bgcolor='#0d1117', plot_bgcolor='#0d1117')
            st.plotly_chart(fig3, use_container_width=True)
        with col_d:
            if len(num_cols) >= 2:
                fig4 = px.box(df, x=target, y=num_cols[0], color=target,
                              title=f"{num_cols[0]} by Churn", template='plotly_dark',
                              color_discrete_sequence=['#38bdf8', '#f87171'])
                fig4.update_layout(paper_bgcolor='#0d1117', plot_bgcolor='#0d1117')
                st.plotly_chart(fig4, use_container_width=True)

    if len(num_cols) > 1:
        st.markdown('<div class="sec">Correlation Heatmap</div>', unsafe_allow_html=True)
        corr = df[num_cols].corr()
        fig5 = px.imshow(corr, template='plotly_dark', color_continuous_scale='Blues',
                         title="Feature Correlation Matrix", text_auto='.2f')
        fig5.update_layout(paper_bgcolor='#0d1117')
        st.plotly_chart(fig5, use_container_width=True)

# ── Tab 3: ML Model ───────────────────────────────────────────
with tab3:
    st.markdown(f'<div class="sec">Train {model_name}</div>', unsafe_allow_html=True)

    if st.button(f"🚀 Train {model_name} Now"):
        with st.spinner(f"Training {model_name}... please wait"):
            r = train_model(df, target, model_name)
            st.session_state.r = r
            st.session_state.trained = True
        st.success(f"✅ {model_name} trained successfully!")

    if st.session_state.get('trained'):
        r = st.session_state.r

        m1, m2, m3, m4 = st.columns(4)
        with m1:
            st.markdown(f'<div class="kpi"><div class="kpi-val">{r["acc"]*100:.1f}%</div><div class="kpi-lbl">Accuracy</div><div class="kpi-tag {"tag-g" if r["acc"]>0.75 else "tag-r"}">{"Good" if r["acc"]>0.75 else "Improve"}</div></div>', unsafe_allow_html=True)
        with m2:
            st.markdown(f'<div class="kpi"><div class="kpi-val">{r["auc"]:.3f}</div><div class="kpi-lbl">ROC-AUC</div><div class="kpi-tag {"tag-g" if r["auc"]>0.70 else "tag-r"}">{"Good" if r["auc"]>0.70 else "Improve"}</div></div>', unsafe_allow_html=True)
        with m3:
            st.markdown(f'<div class="kpi"><div class="kpi-val">{r["prec"]*100:.1f}%</div><div class="kpi-lbl">Precision</div><div class="kpi-tag tag-b">Score</div></div>', unsafe_allow_html=True)
        with m4:
            st.markdown(f'<div class="kpi"><div class="kpi-val">{r["rec"]*100:.1f}%</div><div class="kpi-lbl">Recall</div><div class="kpi-tag tag-b">Score</div></div>', unsafe_allow_html=True)

        st.markdown("<br>", unsafe_allow_html=True)
        col_e, col_f = st.columns(2)

        with col_e:
            fig_cm = px.imshow(r['cm'], text_auto=True, template='plotly_dark',
                               title="Confusion Matrix",
                               labels=dict(x="Predicted", y="Actual"),
                               color_continuous_scale='Blues')
            fig_cm.update_layout(paper_bgcolor='#0d1117')
            st.plotly_chart(fig_cm, use_container_width=True)

        with col_f:
            fig_roc = go.Figure()
            fig_roc.add_trace(go.Scatter(x=r['fpr'], y=r['tpr'], mode='lines',
                                          name=f'AUC={r["auc"]:.3f}',
                                          line=dict(color='#38bdf8', width=2.5)))
            fig_roc.add_trace(go.Scatter(x=[0,1], y=[0,1], mode='lines',
                                          name='Random', line=dict(color='#475569', dash='dash')))
            fig_roc.update_layout(title='ROC Curve', template='plotly_dark',
                                   paper_bgcolor='#0d1117', plot_bgcolor='#0d1117',
                                   xaxis_title='False Positive Rate',
                                   yaxis_title='True Positive Rate')
            st.plotly_chart(fig_roc, use_container_width=True)

        st.markdown('<div class="sec">🏆 Feature Importance</div>', unsafe_allow_html=True)
        fig_fi = px.bar(r['fi'].head(12), x='Importance', y='Feature', orientation='h',
                        title="Top Features Driving Churn", template='plotly_dark',
                        color='Importance', color_continuous_scale='Blues')
        fig_fi.update_layout(paper_bgcolor='#0d1117', plot_bgcolor='#0d1117', height=420)
        st.plotly_chart(fig_fi, use_container_width=True)

    else:
        st.markdown('<div class="card">Click the Train button above to start training the ML model!</div>', unsafe_allow_html=True)

# ── Tab 4: Predict ────────────────────────────────────────────
with tab4:
    st.markdown('<div class="sec">🎯 Predict for a New Customer</div>', unsafe_allow_html=True)

    if not st.session_state.get('trained'):
        st.markdown('<div class="card card-y">⚠️ Please train the model first in the ML Model tab!</div>', unsafe_allow_html=True)
    else:
        r = st.session_state.r
        df_proc = preprocess(df, target)
        X_samp = df_proc.drop(columns=[target])

        st.markdown("Fill in customer details below:")
        input_data = {}
        cols3 = st.columns(3)
        for i, col in enumerate(r['X_cols']):
            with cols3[i % 3]:
                mn = float(X_samp[col].min())
                mx = float(X_samp[col].max())
                mv = float(X_samp[col].mean())
                input_data[col] = st.number_input(col, min_value=mn, max_value=mx, value=mv, key=f"pred_{col}")

        if st.button("🔮 Predict Now"):
            inp = pd.DataFrame([input_data])
            inp_sc = r['scaler'].transform(inp)
            pred = r['model'].predict(inp_sc)[0]
            prob = r['model'].predict_proba(inp_sc)[0]
            churn_p = prob[1] * 100 if r['n_cls'] == 2 else prob.max() * 100

            st.markdown("<br>", unsafe_allow_html=True)
            if pred == 1:
                st.markdown(f'<div class="card card-r">🚨 <strong>HIGH CHURN RISK</strong> — This customer is likely to leave.<br>Churn Probability: <strong>{churn_p:.1f}%</strong><br><br>💡 Recommendation: Offer a discount, upgrade plan, or personal outreach immediately.</div>', unsafe_allow_html=True)
            else:
                st.markdown(f'<div class="card card-g">✅ <strong>LOW CHURN RISK</strong> — This customer is likely to stay.<br>Retention Probability: <strong>{100-churn_p:.1f}%</strong><br><br>💡 Recommendation: Continue current engagement strategy.</div>', unsafe_allow_html=True)

            gauge = go.Figure(go.Indicator(
                mode="gauge+number+delta",
                value=churn_p,
                title={'text': "Churn Risk %", 'font': {'color': '#94a3b8'}},
                delta={'reference': 50, 'increasing': {'color': '#f87171'}, 'decreasing': {'color': '#34d399'}},
                gauge={
                    'axis': {'range': [0, 100], 'tickcolor': '#475569'},
                    'bar': {'color': '#f87171' if churn_p > 50 else '#34d399'},
                    'bgcolor': '#0d1117',
                    'steps': [
                        {'range': [0, 30], 'color': '#052e16'},
                        {'range': [30, 60], 'color': '#1c1917'},
                        {'range': [60, 100], 'color': '#1c0505'}
                    ],
                    'threshold': {'line': {'color': '#fbbf24', 'width': 3}, 'value': 50}
                }
            ))
            gauge.update_layout(paper_bgcolor='#0d1117', font_color='#94a3b8', height=320)
            st.plotly_chart(gauge, use_container_width=True)

# ── Tab 5: Export ─────────────────────────────────────────────
with tab5:
    st.markdown('<div class="sec">📥 Export Your Analysis</div>', unsafe_allow_html=True)

    if not st.session_state.get('trained'):
        st.markdown('<div class="card card-y">⚠️ Please train the model first before exporting!</div>', unsafe_allow_html=True)
    else:
        r = st.session_state.r
        export_name = st.text_input("File name:", value=fname.replace('.csv', ''))

        st.markdown("#### Choose export format:")
        col_x, col_y = st.columns(2)

        with col_x:
            st.markdown("""
            <div class="card">
            <strong>📊 Excel Workbook</strong><br><br>
            Contains 5 sheets:<br>
            • Raw Data (formatted table)<br>
            • Summary Statistics<br>
            • Model Performance metrics<br>
            • Feature Importance ranking<br>
            • Predictions (actual vs predicted)
            </div>
            """, unsafe_allow_html=True)
            excel_data = to_excel(df, r, fname)
            st.download_button(
                "⬇️ Download Excel (.xlsx)",
                data=excel_data,
                file_name=f"{export_name}_churn_analysis.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )

        with col_y:
            st.markdown("""
            <div class="card">
            <strong>📝 Word Document</strong><br><br>
            Contains 5 sections:<br>
            • Model Performance Summary<br>
            • Top Feature Importance<br>
            • Dataset Overview<br>
            • Summary Statistics table<br>
            • Data Sample (first 10 rows)
            </div>
            """, unsafe_allow_html=True)
            word_data = to_word(df, r, fname)
            st.download_button(
                "⬇️ Download Word (.docx)",
                data=word_data,
                file_name=f"{export_name}_churn_report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        st.markdown('<div class="sec">📋 Quick Summary</div>', unsafe_allow_html=True)
        best_feat = r['fi'].iloc[0]['Feature']
        worst_feat = r['fi'].iloc[-1]['Feature']

        insights = [
            ("card-g", f"✅ Model accuracy is <strong>{r['acc']*100:.1f}%</strong> — {'great result!' if r['acc'] > 0.80 else 'solid baseline for a fresher project.'}"),
            ("card-b", f"🏆 Most important churn factor: <strong>{best_feat}</strong> — focus business decisions here."),
            ("card-y", f"📉 Least impactful feature: <strong>{worst_feat}</strong> — consider removing it to simplify the model."),
            ("card-g", f"🎯 ROC-AUC of <strong>{r['auc']:.3f}</strong> — {'excellent discrimination ability.' if r['auc'] > 0.85 else 'good for a starting model.'}"),
            ("card-b", f"💡 Churn rate in this dataset: <strong>{churn_rate}%</strong> — {'class imbalance detected, consider SMOTE.' if churn_rate < 20 or churn_rate > 80 else 'balanced dataset, great for training.'}"),
        ]
        for cls, text in insights:
            st.markdown(f'<div class="card {cls}">{text}</div>', unsafe_allow_html=True)
